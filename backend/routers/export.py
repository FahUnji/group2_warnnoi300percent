"""
Export router — generates downloadable .xlsx and .docx reports from local SQLite bug data.

Endpoints:
  GET /api/export/bugs/xlsx?project_key=KEY
  GET /api/export/bugs/docx?project_key=KEY
  GET /api/export/sprint/xlsx?project_key=KEY&sprint_name=NAME
  GET /api/export/sprint/docx?project_key=KEY&sprint_name=NAME

All endpoints return StreamingResponse (binary file download).
No auth dependency — auth is ambient (OAuth token in SQLite), matching all other routers.
Error shape: {"ok": false, "error": "<code>", "message": "<text>"}
"""
import io
import re
from datetime import date

import matplotlib
matplotlib.use('Agg')  # non-interactive backend required for server-side rendering
import matplotlib.pyplot as plt
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import StreamingResponse
from openpyxl import Workbook

from backend.database import get_db
from backend.dependencies import get_current_user
from backend.services.jira_sync_service import _validate_project_key

router = APIRouter(prefix="/api/export", tags=["export"])

_DONE_KEYWORDS = ["done", "closed", "resolved", "won't fix", "wontfix", "duplicate", "fixed"]
_IN_PROGRESS_KEYWORDS = ["progress", "review", "testing", "development", "in dev"]


def _classify_status(status: str) -> str:
    s = (status or "").lower()
    if any(k in s for k in _DONE_KEYWORDS):
        return "done"
    if any(k in s for k in _IN_PROGRESS_KEYWORDS):
        return "inProgress"
    return "todo"


def _classify_priority(priority: str) -> str:
    p = (priority or "").lower()
    if "critical" in p or "blocker" in p:
        return "critical"
    if "high" in p or "major" in p:
        return "high"
    if "medium" in p or "normal" in p or "minor" in p:
        return "medium"
    return "low"


def _compute_stats(rows: list) -> dict:
    status_counts = {"todo": 0, "inProgress": 0, "done": 0}
    priority_counts = {"critical": 0, "high": 0, "medium": 0, "low": 0}
    for r in rows:
        status_counts[_classify_status(r["status"])] += 1
        priority_counts[_classify_priority(r["priority"])] += 1
    total = len(rows)
    resolved = status_counts["done"]
    return {
        "total": total,
        "open_count": total - resolved,
        "resolved_count": resolved,
        "status_counts": status_counts,
        "priority_counts": priority_counts,
    }


def _make_priority_donut_chart(priority_counts: dict, total: int) -> bytes:
    keys = ["critical", "high", "medium", "low"]
    colors = ["#ba1a1a", "#f57c00", "#fbc02d", "#b2c5ff"]
    values = [priority_counts.get(k, 0) for k in keys]
    if sum(values) == 0:
        values = [1, 0, 0, 0]
        total = 0
    fig, ax = plt.subplots(figsize=(3, 3))
    ax.pie(values, colors=colors, wedgeprops=dict(width=0.5), startangle=90, counterclock=False)
    ax.text(0, 0.1, str(total), ha='center', va='center', fontsize=16, fontweight='bold', color='#dc2626')
    ax.text(0, -0.2, 'Total Bugs', ha='center', va='center', fontsize=7, color='#6b7280')
    buf = io.BytesIO()
    fig.savefig(buf, format="png", bbox_inches="tight", dpi=150)
    plt.close(fig)
    buf.seek(0)
    return buf.read()


def _make_status_donut_chart(status_counts: dict, done_pct: int) -> bytes:
    keys = ["todo", "inProgress", "done"]
    colors = ["#d1d5db", "#f59e0b", "#065b41"]
    values = [status_counts.get(k, 0) for k in keys]
    if sum(values) == 0:
        values = [1, 0, 0]
        done_pct = 0
    fig, ax = plt.subplots(figsize=(3, 3))
    ax.pie(values, colors=colors, wedgeprops=dict(width=0.5), startangle=90, counterclock=False)
    ax.text(0, 0.1, f'{done_pct}%', ha='center', va='center', fontsize=16, fontweight='bold', color='#065b41')
    ax.text(0, -0.2, 'Done', ha='center', va='center', fontsize=7, color='#6b7280')
    buf = io.BytesIO()
    fig.savefig(buf, format="png", bbox_inches="tight", dpi=150)
    plt.close(fig)
    buf.seek(0)
    return buf.read()


def _remove_table_borders(table) -> None:
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'none')
        tblBorders.append(border)
    tblPr.append(tblBorders)


def _set_cell_bg(cell, hex_color: str) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def _slugify(name: str) -> str:
    return re.sub(r"[^a-z0-9]+", "-", name.lower()).strip("-")


def _display_state(state: str) -> str:
    _MAP = {"released": "Completed", "closed": "Completed", "active": "Active",
            "upcoming": "Upcoming", "archived": "Archived"}
    return _MAP.get((state or "").lower(), (state or "unknown").title())


def _build_xlsx_buf(rows: list, headers: list) -> io.BytesIO:
    wb = Workbook()
    ws = wb.active
    ws.title = "Bugs"
    ws.append(headers)
    for cell in ws[1]:
        cell.font = cell.font.copy(bold=True)
    for row in rows:
        ws.append(list(row))
    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    return buf


def _add_word_table(doc: Document, headers: list, rows: list) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        run = hdr_cells[i].paragraphs[0].runs[0]
        run.bold = True
    for row_data in rows:
        row_cells = table.add_row().cells
        for i, val in enumerate(row_data):
            row_cells[i].text = str(val or "")


@router.get("/bugs/xlsx")
async def export_bugs_xlsx(
    project_key: str,
    user_id: int = Depends(get_current_user),
):
    try:
        _validate_project_key(project_key)
    except ValueError as exc:
        raise HTTPException(
            status_code=400,
            detail={"ok": False, "error": "invalid_project_key", "message": str(exc)},
        )
    conn = get_db()
    try:
        rows = conn.execute(
            "SELECT issue_key, summary, status, priority, assignee, sprint_name"
            " FROM bugs WHERE project_key = ? AND user_id = ? ORDER BY issue_key ASC",
            (project_key, user_id),
        ).fetchall()
        proj_row = conn.execute(
            "SELECT project_name FROM jira_projects WHERE project_key = ? AND user_id = ?",
            (project_key, user_id),
        ).fetchone()
    finally:
        conn.close()
    project_name = proj_row["project_name"] if proj_row and proj_row["project_name"] else ""
    stats = _compute_stats(rows)
    pc = stats["priority_counts"]
    total = stats["total"]

    def _pct(n: int) -> str:
        return f"{round(n / total * 100)}%" if total else "0%"

    wb = Workbook()
    ws = wb.active
    ws.title = "Bug Report"

    # --- Summary block ---
    ws.append(["Project", project_key])
    if project_name:
        ws.append(["Project Name", project_name])
    ws.append(["Generated", str(date.today())])
    ws.append([])
    ws.append(["Total Bugs", total])
    ws.append(["Open", stats["open_count"]])
    ws.append(["Resolved", stats["resolved_count"]])
    ws.append([])
    ws.append(["Critical", pc["critical"], _pct(pc["critical"])])
    ws.append(["High", pc["high"], _pct(pc["high"])])
    ws.append(["Medium", pc["medium"], _pct(pc["medium"])])
    ws.append(["Low", pc["low"], _pct(pc["low"])])
    ws.append([])
    for row_cells in ws.iter_rows(min_row=1, max_row=ws.max_row, min_col=1, max_col=1):
        for cell in row_cells:
            if cell.value:
                cell.font = cell.font.copy(bold=True)

    # --- Bug table ---
    headers = ["ID", "Summary", "Status", "Priority", "Assignee", "Fix Version"]
    ws.append(headers)
    for cell in ws[ws.max_row]:
        cell.font = cell.font.copy(bold=True)
    for r in rows:
        ws.append([r["issue_key"], r["summary"], r["status"], r["priority"], r["assignee"], r["sprint_name"]])

    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    filename = f"{project_key}-bug-report-{date.today()}.xlsx"
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


@router.get("/bugs/docx")
async def export_bugs_docx(
    project_key: str,
    user_id: int = Depends(get_current_user),
):
    try:
        _validate_project_key(project_key)
    except ValueError as exc:
        raise HTTPException(
            status_code=400,
            detail={"ok": False, "error": "invalid_project_key", "message": str(exc)},
        )
    conn = get_db()
    try:
        rows = conn.execute(
            "SELECT issue_key, summary, status, priority, assignee, sprint_name"
            " FROM bugs WHERE project_key = ? AND user_id = ? ORDER BY issue_key ASC",
            (project_key, user_id),
        ).fetchall()
        proj_row = conn.execute(
            "SELECT project_name FROM jira_projects WHERE project_key = ? AND user_id = ?",
            (project_key, user_id),
        ).fetchone()
    finally:
        conn.close()
    project_name = proj_row["project_name"] if proj_row and proj_row["project_name"] else ""

    stats = _compute_stats(rows)
    pc = stats["priority_counts"]
    sc = stats["status_counts"]
    total = stats["total"]
    done_pct = round((stats["resolved_count"] / total) * 100) if total > 0 else 0

    def _pct(n: int) -> str:
        return f"{round(n / total * 100)}%" if total else "0%"

    doc = Document()

    # ── Title ──
    title_para = doc.add_paragraph()
    title_run = title_para.add_run(f"{project_key} Bug Report")
    title_run.bold = True
    title_run.font.size = Pt(22)
    title_run.font.color.rgb = RGBColor(0x00, 0x2d, 0x1c)
    title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # ── Metadata ──
    meta_parts = [f"Project: {project_key}"]
    if project_name:
        meta_parts.append(f"  ·  {project_name}")
    meta_parts.append(f"  ·  Generated: {date.today()}")
    meta_para = doc.add_paragraph("".join(meta_parts))
    meta_para.runs[0].font.size = Pt(9)
    meta_para.runs[0].font.color.rgb = RGBColor(0x43, 0x46, 0x54)
    doc.add_paragraph()

    # ── Summary stats table (3 cols) ──
    doc.add_paragraph("Summary").runs[0].bold = True
    stats_table = doc.add_table(rows=2, cols=3)
    _remove_table_borders(stats_table)
    stats_data = [
        ("Total Bugs", str(total), "#dc2626"),
        ("Open", str(stats["open_count"]), "#f59e0b"),
        ("Resolved", f"{stats['resolved_count']}  ({done_pct}%)", "#065b41"),
    ]
    for col_idx, (label, value, _color) in enumerate(stats_data):
        label_cell = stats_table.rows[0].cells[col_idx]
        label_para = label_cell.paragraphs[0]
        label_run = label_para.add_run(label)
        label_run.bold = True
        label_run.font.size = Pt(9)
        label_run.font.color.rgb = RGBColor(0x43, 0x46, 0x54)
        label_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        val_cell = stats_table.rows[1].cells[col_idx]
        val_para = val_cell.paragraphs[0]
        val_run = val_para.add_run(value)
        val_run.bold = True
        val_run.font.size = Pt(20)
        val_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    # ── Charts side by side ──
    doc.add_paragraph("Bug Analysis").runs[0].bold = True
    priority_chart_bytes = _make_priority_donut_chart(pc, total)
    status_chart_bytes = _make_status_donut_chart(sc, done_pct)

    chart_table = doc.add_table(rows=1, cols=2)
    _remove_table_borders(chart_table)

    # Left: priority chart + legend
    left_cell = chart_table.rows[0].cells[0]
    left_para = left_cell.paragraphs[0]
    left_title = left_para.add_run("Bugs by Priority\n")
    left_title.bold = True
    left_title.font.size = Pt(10)
    left_img_run = left_para.add_run()
    left_img_run.add_picture(io.BytesIO(priority_chart_bytes), width=Inches(2.6))
    left_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p_legend_rows = [
        ("Critical", pc["critical"], _pct(pc["critical"])),
        ("High",     pc["high"],     _pct(pc["high"])),
        ("Medium",   pc["medium"],   _pct(pc["medium"])),
        ("Low",      pc["low"],      _pct(pc["low"])),
    ]
    p_legend = left_cell.add_table(rows=len(p_legend_rows), cols=3)
    _remove_table_borders(p_legend)
    for r_idx, (lbl, cnt, pct_str) in enumerate(p_legend_rows):
        p_legend.rows[r_idx].cells[0].paragraphs[0].add_run(lbl).font.size = Pt(8)
        cnt_run = p_legend.rows[r_idx].cells[1].paragraphs[0].add_run(str(cnt))
        cnt_run.bold = True
        cnt_run.font.size = Pt(8)
        p_legend.rows[r_idx].cells[2].paragraphs[0].add_run(pct_str).font.size = Pt(8)

    # Right: status chart + legend
    right_cell = chart_table.rows[0].cells[1]
    right_para = right_cell.paragraphs[0]
    right_title = right_para.add_run("Status Distribution\n")
    right_title.bold = True
    right_title.font.size = Pt(10)
    right_img_run = right_para.add_run()
    right_img_run.add_picture(io.BytesIO(status_chart_bytes), width=Inches(2.6))
    right_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    s_legend_rows = [
        ("To Do",       sc["todo"],       _pct(sc["todo"])),
        ("In Progress", sc["inProgress"], _pct(sc["inProgress"])),
        ("Done",        sc["done"],       _pct(sc["done"])),
    ]
    s_legend = right_cell.add_table(rows=len(s_legend_rows), cols=3)
    _remove_table_borders(s_legend)
    for r_idx, (lbl, cnt, pct_str) in enumerate(s_legend_rows):
        s_legend.rows[r_idx].cells[0].paragraphs[0].add_run(lbl).font.size = Pt(8)
        cnt_run = s_legend.rows[r_idx].cells[1].paragraphs[0].add_run(str(cnt))
        cnt_run.bold = True
        cnt_run.font.size = Pt(8)
        s_legend.rows[r_idx].cells[2].paragraphs[0].add_run(pct_str).font.size = Pt(8)

    doc.add_paragraph()

    # ── Bug list table ──
    doc.add_paragraph("Bug List").runs[0].bold = True
    headers = ["ID", "Summary", "Status", "Priority", "Assignee", "Fix Version"]
    data_rows = [
        [r["issue_key"], r["summary"], r["status"], r["priority"], r["assignee"], r["sprint_name"]]
        for r in rows
    ]
    _add_word_table(doc, headers, data_rows)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    filename = f"{project_key}-bug-report-{date.today()}.docx"
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


@router.get("/sprint/xlsx")
async def export_sprint_xlsx(
    project_key: str,
    sprint_name: str = "",
    user_id: int = Depends(get_current_user),
):
    try:
        _validate_project_key(project_key)
    except ValueError as exc:
        raise HTTPException(
            status_code=400,
            detail={"ok": False, "error": "invalid_project_key", "message": str(exc)},
        )
    conn = get_db()
    try:
        if sprint_name:
            rows = conn.execute(
                "SELECT issue_key, summary, status, priority, assignee, sprint_name"
                " FROM bugs WHERE project_key = ? AND user_id = ? AND sprint_name = ? ORDER BY issue_key ASC",
                (project_key, user_id, sprint_name),
            ).fetchall()
            sprints_meta_rows = conn.execute(
                "SELECT sprint_name, state, start_date, end_date"
                " FROM sprints WHERE project_key = ? AND user_id = ? AND sprint_name = ? LIMIT 1",
                (project_key, user_id, sprint_name),
            ).fetchall()
        else:
            sprints_meta_rows = conn.execute(
                "SELECT sprint_name, state, start_date, end_date"
                " FROM sprints WHERE project_key = ? AND user_id = ? ORDER BY sprint_id DESC",
                (project_key, user_id),
            ).fetchall()
            sprint_names = [r["sprint_name"] for r in sprints_meta_rows]
            rows = conn.execute(
                "SELECT issue_key, summary, status, priority, assignee, sprint_name"
                " FROM bugs WHERE project_key = ? AND user_id = ? ORDER BY sprint_name, issue_key ASC",
                (project_key, user_id),
            ).fetchall()
    finally:
        conn.close()

    # sprint metadata lookup: {sprint_name -> row}
    sprint_meta_by_name = {r["sprint_name"]: r for r in sprints_meta_rows}

    headers = ["ID", "Summary", "Status", "Priority", "Assignee", "Fix Version"]

    if sprint_name:
        data_rows = [
            [r["issue_key"], r["summary"], r["status"], r["priority"], r["assignee"], r["sprint_name"]]
            for r in rows
        ]
        buf = _build_xlsx_buf(data_rows, headers)
        filename = f"{project_key}-{_slugify(sprint_name)}-{date.today()}.xlsx"
    else:
        # One sheet per sprint — order by sprint_id DESC, then unrecognised, then No Sprint
        bugs_by_sprint: dict = {}
        for bug in rows:
            key = bug["sprint_name"] or "No Sprint"
            bugs_by_sprint.setdefault(key, []).append(bug)

        known_set = set(sprint_names)
        unknown = [k for k in bugs_by_sprint if k not in known_set and k != "No Sprint"]
        sheet_order = sprint_names + sorted(unknown) + (["No Sprint"] if "No Sprint" in bugs_by_sprint else [])

        if not sheet_order:
            sheet_order = ["No Data"]
            bugs_by_sprint["No Data"] = []

        wb = Workbook()
        wb.remove(wb.active)  # remove default empty sheet
        for sn in sheet_order:
            sheet_title = re.sub(r'[\\/*?:\[\]]', '-', sn)[:31]
            ws = wb.create_sheet(title=sheet_title)

            # --- Sprint metadata block ---
            meta = sprint_meta_by_name.get(sn)
            ws.append(["Sprint", sn])
            ws.append(["Status", _display_state(meta["state"] if meta else "")])
            ws.append(["Start Date", meta["start_date"] or "N/A" if meta else "N/A"])
            ws.append(["Release Date", meta["end_date"] or "N/A" if meta else "N/A"])
            for row_cells in ws.iter_rows(min_row=1, max_row=4, min_col=1, max_col=1):
                for cell in row_cells:
                    cell.font = cell.font.copy(bold=True)
            ws.append([])  # blank row

            # --- Stats block (always shown, even if all zeros) ---
            sprint_bugs = bugs_by_sprint.get(sn, [])
            stats = _compute_stats(sprint_bugs)
            pc = stats["priority_counts"]
            ws.append(["Total Bugs", stats["total"]])
            ws.append(["Open", stats["open_count"]])
            ws.append(["Resolved", stats["resolved_count"]])
            ws.append(["Critical", pc["critical"]])
            ws.append(["High", pc["high"]])
            ws.append(["Medium", pc["medium"]])
            ws.append(["Low", pc["low"]])
            for row_cells in ws.iter_rows(min_row=6, max_row=12, min_col=1, max_col=1):
                for cell in row_cells:
                    cell.font = cell.font.copy(bold=True)
            ws.append([])  # blank row

            # --- Bug table ---
            ws.append(headers)
            for cell in ws[ws.max_row]:
                cell.font = cell.font.copy(bold=True)
            for bug in sprint_bugs:
                ws.append([bug["issue_key"], bug["summary"], bug["status"],
                            bug["priority"], bug["assignee"], bug["sprint_name"]])

        buf = io.BytesIO()
        wb.save(buf)
        buf.seek(0)
        filename = f"{project_key}-all-sprints-{date.today()}.xlsx"

    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


def _add_sprint_section(doc: Document, sprint_name: str, sprint_meta, bugs: list) -> None:
    stats = _compute_stats(bugs)
    pc = stats["priority_counts"]
    total = stats["total"]
    done_pct = round((stats["resolved_count"] / total) * 100) if total > 0 else 0

    state_label = _display_state(sprint_meta["state"] if sprint_meta else "")
    start_date = (sprint_meta["start_date"] or "N/A") if sprint_meta else "N/A"
    end_date = (sprint_meta["end_date"] or "N/A") if sprint_meta else "N/A"

    _STATE_BG = {"Active": "065b41", "Completed": "374151", "Upcoming": "1d4ed8", "Archived": "6b7280"}
    state_bg = _STATE_BG.get(state_label, "6b7280")

    def _hex_rgb(h: str):
        return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))

    # ── Sprint card (bordered table) ──
    card = doc.add_table(rows=7, cols=4)
    card.style = "Table Grid"

    # Row 0: sprint name (cols 0-2 merged, dark green) | state badge (col 3, state color)
    name_cell = card.rows[0].cells[0]
    name_cell.merge(card.rows[0].cells[2])
    _set_cell_bg(name_cell, "1b4332")
    name_run = name_cell.paragraphs[0].add_run(sprint_name)
    name_run.bold = True
    name_run.font.size = Pt(12)
    name_run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    state_cell = card.rows[0].cells[3]
    _set_cell_bg(state_cell, state_bg)
    state_run = state_cell.paragraphs[0].add_run(state_label.upper())
    state_run.bold = True
    state_run.font.size = Pt(9)
    state_run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    state_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Row 1: date range (all 4 merged, light green-gray)
    date_cell = card.rows[1].cells[0]
    date_cell.merge(card.rows[1].cells[3])
    _set_cell_bg(date_cell, "f0f4f1")
    date_run = date_cell.paragraphs[0].add_run(f"{start_date}  →  {end_date}")
    date_run.font.size = Pt(9)
    date_run.font.color.rgb = RGBColor(0x43, 0x46, 0x54)

    # Row 2: stat labels (light gray bg)
    for col_idx, label in enumerate(["Total Bugs", "Resolved", "Open", "Progress"]):
        cell = card.rows[2].cells[col_idx]
        _set_cell_bg(cell, "f3f4f6")
        run = cell.paragraphs[0].add_run(label)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x43, 0x46, 0x54)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Row 3: stat values (colored text)
    for col_idx, (value, color_hex) in enumerate(zip(
        [str(total), f"{stats['resolved_count']} ({done_pct}%)", str(stats["open_count"]), f"{done_pct}%"],
        ["dc2626", "065b41", "f59e0b", "065b41"],
    )):
        cell = card.rows[3].cells[col_idx]
        _set_cell_bg(cell, "ffffff")
        run = cell.paragraphs[0].add_run(value)
        run.bold = True
        run.font.size = Pt(14)
        run.font.color.rgb = _hex_rgb(color_hex)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Row 4: severity header (all 4 merged, light green)
    sev_hdr_cell = card.rows[4].cells[0]
    sev_hdr_cell.merge(card.rows[4].cells[3])
    _set_cell_bg(sev_hdr_cell, "d1fae5")
    sev_run = sev_hdr_cell.paragraphs[0].add_run("Bug Severity Distribution")
    sev_run.bold = True
    sev_run.font.size = Pt(9)
    sev_run.font.color.rgb = RGBColor(0x06, 0x5b, 0x41)

    # Row 5: severity labels (tinted bg per severity)
    for col_idx, (label, bg) in enumerate(zip(
        ["Critical", "High", "Medium", "Low"],
        ["fef2f2", "fff7ed", "fefce8", "eff6ff"],
    )):
        cell = card.rows[5].cells[col_idx]
        _set_cell_bg(cell, bg)
        run = cell.paragraphs[0].add_run(label)
        run.bold = True
        run.font.size = Pt(9)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Row 6: severity values (colored text matching severity)
    for col_idx, (value, color_hex) in enumerate(zip(
        [str(pc["critical"]), str(pc["high"]), str(pc["medium"]), str(pc["low"])],
        ["dc2626", "f57c00", "eab308", "6366f1"],
    )):
        cell = card.rows[6].cells[col_idx]
        run = cell.paragraphs[0].add_run(value)
        run.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = _hex_rgb(color_hex)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # ── Bug list ──
    bug_list_para = doc.add_paragraph("Bug List")
    if bug_list_para.runs:
        bug_list_para.runs[0].bold = True
    if bugs:
        _add_word_table(doc, ["ID", "Summary", "Status", "Priority", "Assignee", "Fix Version"], [
            [r["issue_key"], r["summary"], r["status"], r["priority"], r["assignee"], r["sprint_name"]]
            for r in bugs
        ])
    else:
        doc.add_paragraph("No bugs recorded for this sprint.")
    doc.add_paragraph()


@router.get("/sprint/docx")
async def export_sprint_docx(
    project_key: str,
    sprint_name: str = "",
    user_id: int = Depends(get_current_user),
):
    try:
        _validate_project_key(project_key)
    except ValueError as exc:
        raise HTTPException(
            status_code=400,
            detail={"ok": False, "error": "invalid_project_key", "message": str(exc)},
        )
    conn = get_db()
    try:
        proj_row = conn.execute(
            "SELECT project_name FROM jira_projects WHERE project_key = ? AND user_id = ?",
            (project_key, user_id),
        ).fetchone()
        if sprint_name:
            sprints_meta = conn.execute(
                "SELECT sprint_id, sprint_name, state, start_date, end_date"
                " FROM sprints WHERE project_key = ? AND user_id = ? AND sprint_name = ? LIMIT 1",
                (project_key, user_id, sprint_name),
            ).fetchall()
            all_bugs = conn.execute(
                "SELECT issue_key, summary, status, priority, assignee, sprint_name"
                " FROM bugs WHERE project_key = ? AND user_id = ? AND sprint_name = ? ORDER BY issue_key ASC",
                (project_key, user_id, sprint_name),
            ).fetchall()
        else:
            sprints_meta = conn.execute(
                "SELECT sprint_id, sprint_name, state, start_date, end_date"
                " FROM sprints WHERE project_key = ? AND user_id = ? ORDER BY sprint_id DESC",
                (project_key, user_id),
            ).fetchall()
            all_bugs = conn.execute(
                "SELECT issue_key, summary, status, priority, assignee, sprint_name"
                " FROM bugs WHERE project_key = ? AND user_id = ? ORDER BY sprint_name, issue_key ASC",
                (project_key, user_id),
            ).fetchall()
    finally:
        conn.close()

    project_name = proj_row["project_name"] if proj_row and proj_row["project_name"] else ""
    doc = Document()

    def _add_doc_title(text: str) -> None:
        p = doc.add_paragraph()
        r = p.add_run(text)
        r.bold = True
        r.font.size = Pt(22)
        r.font.color.rgb = RGBColor(0x00, 0x2d, 0x1c)

    def _add_doc_meta(extra: str = "") -> None:
        parts = [f"Project: {project_key}"]
        if project_name:
            parts.append(f"  ·  {project_name}")
        parts.append(f"  ·  Generated: {date.today()}")
        if extra:
            parts.append(f"  ·  {extra}")
        p = doc.add_paragraph("".join(parts))
        if p.runs:
            p.runs[0].font.size = Pt(9)
            p.runs[0].font.color.rgb = RGBColor(0x43, 0x46, 0x54)

    if sprint_name:
        sprint_meta = sprints_meta[0] if sprints_meta else None
        _add_doc_title(f"{project_key} - {sprint_name} Sprint Report")
        _add_doc_meta()
        doc.add_paragraph()
        _add_sprint_section(doc, sprint_name, sprint_meta, all_bugs)
        filename = f"{project_key}-{_slugify(sprint_name)}-{date.today()}.docx"
    else:
        bugs_by_sprint: dict = {}
        for bug in all_bugs:
            key = bug["sprint_name"] or "No Sprint"
            bugs_by_sprint.setdefault(key, []).append(bug)

        overall_stats = _compute_stats(all_bugs)
        _add_doc_title(f"{project_key} - All Sprints Report")
        _add_doc_meta(f"Total Sprints: {len(sprints_meta)}  ·  Total Bugs: {overall_stats['total']}")
        doc.add_paragraph()

        for sm in sprints_meta:
            sn = sm["sprint_name"]
            _add_sprint_section(doc, sn, sm, bugs_by_sprint.get(sn, []))

        if "No Sprint" in bugs_by_sprint:
            _add_sprint_section(doc, "No Sprint", None, bugs_by_sprint["No Sprint"])

        filename = f"{project_key}-all-sprints-{date.today()}.docx"

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )
